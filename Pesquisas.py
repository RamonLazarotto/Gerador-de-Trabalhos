from cgitb import text
from ctypes import alignment
from pickle import TRUE
from posixpath import split
from pydoc import Doc
import wikipedia
import re
from docx import Document

name = input('Digite o seu nome: ')
wikipedia.set_lang('pt')
tittle = input('Sobre o que você quer pesquisar ? \n')
while TRUE:
    try:
        wiki = wikipedia.page(tittle)
        break
    except:
        print('Nome do projeto inválido')

        tittle = input('Digite outro nome de projeto: \n')
text = wiki.content
text = re.sub(r'==','', text)
text = re.sub(r'=','', text) 
text = re.sub(r'\n','\n', text)
split = text.split('Veja também',1)
text = split[0]

print (text)

Document = Document()
paragraph = Document.add_heading(tittle,0)
paragraph.alignment = 1


paragraph = Document.add_paragraph('  ' + text)
paragraph = Document.add_paragraph(name)
paragraph.alignment = 2
Document.savetittle('.docx')
input()